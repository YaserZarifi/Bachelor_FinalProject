# -*- coding: utf-8 -*-
"""
build_docx.py — رندر گزارش پایان‌نامه از مارک‌داون به Word (.docx)

این اسکریپت خروجی را دقیقاً مطابق «قالب و شیوه‌نامه‌ی نگارش پایان‌نامه یا
رساله‌ی دانشگاه صنعتی امیرکبیر» تولید می‌کند. تناظر بندهای شیوه‌نامه با
پیاده‌سازی:

  بخش ۴-۱ / ۴-۲  قلم‌ها  → B Nazanin برای فارسی، Times New Roman برای لاتین
                          با اندازه‌ی «یک شماره کوچک‌تر»؛ کد با Courier New 10.
                          (در OOXML: w:szCs = اندازه‌ی فارسی، w:sz = اندازه‌ی لاتین)
  بخش ۴-۴-۱     حاشیه‌ها → بالا ۳، پایین ۳، راست ۳، چپ ۲/۵ سانتی‌متر
  بخش ۴-۴-۲     فاصله‌ی خطوط ۱/۵
  بخش ۴-۴-۳     فاصله‌ی پیش از بند ۶pt و پیش از عنوان‌ها ۳۰/۲۷/۲۰/۱۵pt
  بخش ۴-۸/۵-۱۲  سربرگ: عنوان فصل در سمت راست + خط پررنگ زیر آن؛
                          ته‌برگ: شماره‌ی صفحه در پایین و وسط
  بخش ۴-۹       شماره‌گذاری فصل-ترتیب؛ عنوان جدول بالا، عنوان شکل پایین
  بخش ۳-۱       پاورقیِ معادل لاتین در نخستین ظهور اصطلاح/اختصار
  بخش ۵-۵       بدون تورفتگیِ خودکارِ ابتدای بند
  پیوست         شماره‌گذاری پ-۱ برای شکل/جدولِ پیوست
  صفحات مقدماتی شماره‌گذاری أ، ب، ت، ث ... (arabicAlpha)

اجرا:  python thesis/build_docx.py
خروجی: thesis/out/report.docx  +  گزارش اعتبارسنجی در کنسول.

مارک‌داون پشتیبانی‌شده (زیرمجموعه‌ی سفارشی):
  # عنوان                  → عنوان فصل (صفحه‌ی جداکننده + سربرگ اختصاصی)
  ## / ### / ####          → بخش/زیربخش/زیرزیربخش (شماره‌گذاری خودکارِ فصل-ترتیب)
  متن ساده                 → بند فارسیِ دوطرفه‌چین
  ```lang ... ```          → بلوک کد (چپ‌چین، Courier New)
  - مورد / 1. مورد         → فهرست نقطه‌ای/عددی
  | a | b |                → جدول (سطر اول = سرستون)
  [FIGURE file=.. | label=.. | caption=..]   → شکل + عنوان زیر آن
  [TABLE label=.. | caption=..]              → عنوان جدولِ بعدی (بالای جدول)
  متن^[پاورقی]             → پاورقی (معادل لاتین اصطلاح)
  > نقل‌قول                → بند تورفته
  [[SCREENSHOT-XX]]        → جای اسکرین‌شات
  **bold**  و  `code`      → درون‌خطی
  {{PAGEBREAK}}            → شکست صفحه
"""

from __future__ import annotations
import os
import re
from xml.sax.saxutils import escape as xml_escape

try:
    import yaml
except Exception:
    yaml = None

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.opc.part import Part
from docx.opc.packuri import PackURI

# ---------------------------------------------------------------------------
# مسیرها و پیکربندی
# ---------------------------------------------------------------------------
HERE = os.path.dirname(os.path.abspath(__file__))
CONTENT_DIR = os.path.join(HERE, "content")
FIG_DIR = os.path.join(HERE, "figures")
OUT_DIR = os.path.join(HERE, "out")
OUT_FILE = os.path.join(OUT_DIR, "report.docx")
CONFIG_FILE = os.path.join(HERE, "config.yaml")

DEFAULT_CFG = {
    "meta": {},
    "fonts": {
        "body_fa": "B Nazanin", "body_latin": "Times New Roman",
        "code": "Courier New", "code_size_pt": 10,
        "code_comment_fa": "B Nazanin", "code_comment_size_pt": 11,
    },
    "sizes_fa": {"body": 14, "h1": 20, "h2": 18, "h3": 16, "h4": 14, "h5": 13,
                 "caption": 13, "in_table": 13, "reference": 13, "footnote": 11},
    "latin_size_delta": 1,
    "margins_cm": {"top": 3.0, "bottom": 3.0, "right": 3.0, "left": 2.5,
                   "header_cm": 1.5, "footer_cm": 1.5},
    "line_spacing": 1.5,
    "paragraph_space_before_pt": 6,
    "paragraph_indent_cm": 0,
    "heading_space_before_pt": {"chapter": 30, "h2": 27, "h3": 20, "h4": 15, "h5": 15},
    "toc": {"title": "فهرست عنوان‌ها", "figures_title": "فهرست شکل‌ها",
            "tables_title": "فهرست جدول‌ها", "symbols_title": "فهرست نمادها", "levels": 3},
    "page_numbers": {"front_matter": "arabicAlpha", "body": "decimal"},
    "caption_separator": " ",
    "appendix_prefix": "پ",
    "figure_box": True,
}


def load_cfg():
    cfg = {k: (dict(v) if isinstance(v, dict) else v) for k, v in DEFAULT_CFG.items()}
    if yaml and os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, encoding="utf-8") as f:
            loaded = yaml.safe_load(f) or {}
        for k, v in loaded.items():
            if isinstance(v, dict) and isinstance(cfg.get(k), dict):
                merged = dict(cfg[k]); merged.update(v); cfg[k] = merged
            else:
                cfg[k] = v
    return cfg


# ---------------------------------------------------------------------------
# ارقام فارسی
# ---------------------------------------------------------------------------
_FA_DIGITS = str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹")
CHAPTER_WORDS = ["", "یک", "دو", "سه", "چهار", "پنج", "شش", "هفت", "هشت", "نه", "ده"]


def fa_num(s) -> str:
    return str(s).translate(_FA_DIGITS)


_PERSIAN_RE = re.compile(r"[؀-ۿ]")


def has_persian(text: str) -> bool:
    return bool(_PERSIAN_RE.search(text or ""))


# ---------------------------------------------------------------------------
# کمک‌کارهای سطح‌پایین OOXML
# ---------------------------------------------------------------------------
def _oxml(tag, **attrs):
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k), v)
    return el


def set_par_bidi(paragraph):
    """<w:bidi/> در pPr — پیش از spacing/ind/jc درج می‌شود تا ترتیب اسکیمای OOXML حفظ شود."""
    pPr = paragraph._p.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is not None:
        return
    bidi = _oxml("w:bidi", **{"w:val": "1"})
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is not None:
        pStyle.addnext(bidi)
    else:
        pPr.insert(0, bidi)


def _apply_fonts(rPr, cfg, *, fa_size, mono=False):
    """
    قاعده‌ی جدول ۴-۱/۴-۲ شیوه‌نامه در سطح OOXML:
      w:cs   (Complex Script) = B Nazanin        با اندازه‌ی w:szCs = fa_size
      w:ascii/w:hAnsi (لاتین) = Times New Roman  با اندازه‌ی w:sz   = fa_size − ۱
    این‌گونه، حتی اگر واژه‌ای لاتین داخل یک run فارسی بیفتد، خودبه‌خود با قلم و
    اندازه‌ی درستِ لاتین رندر می‌شود.
    """
    fonts = cfg["fonts"]
    delta = cfg.get("latin_size_delta", 1)
    if mono:
        fa_name = latin_name = fonts["code"]
        fa_size = latin_size = fonts["code_size_pt"]
    else:
        fa_name = fonts["body_fa"]
        latin_name = fonts["body_latin"]
        latin_size = max(fa_size - delta, 6)

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = _oxml("w:rFonts")
        rPr.insert(0, rFonts)
    # سبک‌های آماده‌ی Word (به‌ویژه Heading 1..5) قلم را با ارجاع به «تم» تعریف
    # می‌کنند و طبق اسکیمای OOXML، w:cstheme بر w:cs اولویت دارد. اگر این
    # ارجاع‌ها حذف نشوند، عنوان‌ها به‌جای B Nazanin با قلم تم رندر می‌شوند.
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme"):
        if rFonts.get(qn(attr)) is not None:
            del rFonts.attrib[qn(attr)]
    rFonts.set(qn("w:ascii"), latin_name)
    rFonts.set(qn("w:hAnsi"), latin_name)
    rFonts.set(qn("w:cs"), fa_name)

    for tag, size in (("w:sz", latin_size), ("w:szCs", fa_size)):
        el = rPr.find(qn(tag))
        if el is None:
            el = _oxml(tag)
            rPr.append(el)
        el.set(qn("w:val"), str(int(round(size * 2))))


def style_run(run, cfg, *, rtl=True, mono=False, bold=False, size_pt=None, color=None):
    """اعمال قلم/اندازه/جهت روی یک run طبق قواعد Complex Script."""
    fa_size = size_pt or cfg["sizes_fa"]["body"]
    rPr = run._r.get_or_add_rPr()
    _apply_fonts(rPr, cfg, fa_size=fa_size, mono=mono)

    if bold:
        run.bold = True
        rPr.append(_oxml("w:bCs", **{"w:val": "1"}))   # bold برای Complex Script

    if color is not None:
        run.font.color.rgb = color

    rPr.append(_oxml("w:rtl", **{"w:val": "1" if rtl else "0"}))
    return run


def set_section_rtl(section):
    sectPr = section._sectPr
    if sectPr.find(qn("w:bidi")) is None:
        sectPr.append(_oxml("w:bidi", **{"w:val": "1"}))


def _add_bottom_border(paragraph, size=12, color="000000"):
    """خط پررنگ زیر سربرگ (بخش ۴-۸ شیوه‌نامه)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = _oxml("w:pBdr")
    bottom = _oxml("w:bottom", **{"w:val": "single", "w:sz": str(size),
                                  "w:space": "1", "w:color": color})
    pBdr.append(bottom)
    pPr.append(pBdr)


def _box_paragraph(paragraph, size=6, color="808080"):
    """کادر بیرونی دور شکل (توصیه‌ی بخش «نکته‌ها و تذکرها»)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = _oxml("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        pBdr.append(_oxml(f"w:{edge}", **{"w:val": "single", "w:sz": str(size),
                                          "w:space": "4", "w:color": color}))
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# پاورقی‌ها (بخش ۳-۱ شیوه‌نامه) — تزریق مستقیم بخش footnotes.xml به بسته
# ---------------------------------------------------------------------------
FOOTNOTES_CT = ("application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.footnotes+xml")
FOOTNOTES_RT = ("http://schemas.openxmlformats.org/officeDocument"
                "/2006/relationships/footnotes")
_W_NS = nsmap["w"]


class Footnotes:
    """انباره‌ی پاورقی‌ها؛ در پایانِ ساخت به‌صورت یک part به سند افزوده می‌شود."""

    def __init__(self, cfg):
        self.cfg = cfg
        self.items = []          # [(id, text)]
        self.seen = set()        # جلوگیری از تکرار پاورقی برای یک اصطلاح
        self._next_id = 1

    def add_reference(self, paragraph, text):
        """یک نشانه‌ی پاورقی به پاراگراف می‌افزاید و متن آن را ثبت می‌کند."""
        text = text.strip()
        if not text:
            return None
        fid = self._next_id
        self._next_id += 1
        self.items.append((fid, text))

        run = paragraph.add_run()
        rPr = run._r.get_or_add_rPr()
        _apply_fonts(rPr, self.cfg, fa_size=self.cfg["sizes_fa"]["footnote"])
        rPr.append(_oxml("w:vertAlign", **{"w:val": "superscript"}))
        rPr.append(_oxml("w:rtl", **{"w:val": "0"}))
        run._r.append(_oxml("w:footnoteReference", **{"w:id": str(fid)}))
        return fid

    def once(self, key):
        """True اگر این اصطلاح تا کنون پاورقی نشده باشد."""
        if key in self.seen:
            return False
        self.seen.add(key)
        return True

    # -- ساخت XML ----------------------------------------------------------
    def _para_xml(self, fid, text):
        fa_size = self.cfg["sizes_fa"]["footnote"]
        latin = max(fa_size - self.cfg.get("latin_size_delta", 1), 6)
        fonts = self.cfg["fonts"]
        rtl = "1" if has_persian(text) else "0"
        jc = "right" if has_persian(text) else "left"
        bidi = "<w:bidi/>" if has_persian(text) else ""
        rpr = (f'<w:rPr><w:rFonts w:ascii="{fonts["body_latin"]}" '
               f'w:hAnsi="{fonts["body_latin"]}" w:cs="{fonts["body_fa"]}"/>'
               f'<w:sz w:val="{int(latin*2)}"/><w:szCs w:val="{int(fa_size*2)}"/>'
               f'<w:rtl w:val="{rtl}"/></w:rPr>')
        return (
            f'<w:footnote w:id="{fid}"><w:p><w:pPr>{bidi}'
            f'<w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>'
            f'<w:jc w:val="{jc}"/></w:pPr>'
            f'<w:r><w:rPr><w:rFonts w:ascii="{fonts["body_latin"]}" '
            f'w:hAnsi="{fonts["body_latin"]}" w:cs="{fonts["body_fa"]}"/>'
            f'<w:sz w:val="{int(latin*2)}"/><w:szCs w:val="{int(fa_size*2)}"/>'
            f'<w:vertAlign w:val="superscript"/></w:rPr><w:footnoteRef/></w:r>'
            f'<w:r>{rpr}<w:t xml:space="preserve"> {xml_escape(text)}</w:t></w:r>'
            f'</w:p></w:footnote>'
        )

    def to_blob(self):
        sep = (
            '<w:footnote w:type="separator" w:id="-1"><w:p><w:pPr>'
            '<w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
            '<w:r><w:separator/></w:r></w:p></w:footnote>'
            '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:pPr>'
            '<w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
            '<w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
        )
        body = "".join(self._para_xml(fid, txt) for fid, txt in self.items)
        xml = (f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
               f'<w:footnotes xmlns:w="{_W_NS}">{sep}{body}</w:footnotes>')
        return xml.encode("utf-8")

    def attach(self, doc):
        if not self.items:
            return
        partname = PackURI("/word/footnotes.xml")
        blob = self.to_blob()
        package = doc.part.package
        try:
            part = Part(partname, FOOTNOTES_CT, blob, package)
        except TypeError:                      # امضای نسخه‌های جدیدتر python-docx
            part = Part(partname, FOOTNOTES_CT, package, blob)
        doc.part.relate_to(part, FOOTNOTES_RT)


# ---------------------------------------------------------------------------
# پارس درون‌خطی: تقسیم متن فارسی/لاتین + **bold** + `code` + پاورقی
# ---------------------------------------------------------------------------
# ویرگول و آپاستروف لاتین هم داخل همان run لاتین می‌مانند؛ در غیر این صورت
# عبارتی مانند «Find It, Fix It» در بافت راست‌به‌چپ به دو تکه می‌شکند و جای
# دو نیمه‌ی آن با هم عوض می‌شود.
_LATIN_RE = re.compile(
    r"\(?[A-Za-z0-9][A-Za-z0-9 _\-\+\./:%#&,'’\(\)]*[A-Za-z0-9\)]|[A-Za-z0-9]")
RLM = "‏"  # Right-to-Left Mark

TOKEN_RE = re.compile(
    r"(\^\[[^\]]+\])"                      # پاورقی:  متن^[معادل لاتین]
    r"|(\*\*.+?\*\*)"                      # bold
    r"|(`[^`]+`)"                          # inline code
    r"|(\[\[SCREENSHOT-[0-9A-Za-z]+\]\])"  # screenshot marker
    r"|(\[\[(?:REF|DATA)-[0-9A-Za-z]+:[^\]]*\]\])"  # ref/data marker
)


def add_inline(paragraph, text, cfg, base_rtl=True, size_pt=None, bold_all=False,
               footnotes=None):
    """یک رشته را به runهای فارسی/لاتین/کد/پاورقی تبدیل می‌کند."""
    pos = 0
    for m in TOKEN_RE.finditer(text):
        if m.start() > pos:
            _add_mixed(paragraph, text[pos:m.start()], cfg, base_rtl, size_pt, bold_all)
        tok = m.group(0)
        if tok.startswith("^["):
            if footnotes is not None:
                footnotes.add_reference(paragraph, tok[2:-1])
        elif tok.startswith("**"):
            r = paragraph.add_run(tok[2:-2])
            style_run(r, cfg, rtl=base_rtl, bold=True, size_pt=size_pt)
        elif tok.startswith("`"):
            r = paragraph.add_run(tok[1:-1])
            style_run(r, cfg, rtl=False, mono=True)
            paragraph.add_run(RLM)
        elif tok.startswith("[[SCREENSHOT"):
            r = paragraph.add_run(tok)
            style_run(r, cfg, rtl=False, bold=True, color=RGBColor(0xC0, 0x00, 0x00))
            _highlight(r, "yellow")
        else:  # REF / DATA
            r = paragraph.add_run(tok)
            style_run(r, cfg, rtl=base_rtl, bold=True, color=RGBColor(0xC0, 0x00, 0x00))
        pos = m.end()
    if pos < len(text):
        _add_mixed(paragraph, text[pos:], cfg, base_rtl, size_pt, bold_all)


def _add_mixed(paragraph, text, cfg, base_rtl, size_pt=None, bold_all=False):
    """متن ساده را به قطعات فارسی و لاتین می‌شکند (لاتین → run با rtl=0)."""
    if not text:
        return
    idx = 0
    for m in _LATIN_RE.finditer(text):
        if m.start() > idx:
            r = paragraph.add_run(text[idx:m.start()])
            style_run(r, cfg, rtl=base_rtl, size_pt=size_pt, bold=bold_all)
        r = paragraph.add_run(m.group(0))
        style_run(r, cfg, rtl=False, size_pt=size_pt, bold=bold_all)
        if base_rtl:
            paragraph.add_run(RLM)   # تا نقطه/پرانتزِ پس از عبارت لاتین سر جای خود بماند
        idx = m.end()
    if idx < len(text):
        r = paragraph.add_run(text[idx:])
        style_run(r, cfg, rtl=base_rtl, size_pt=size_pt, bold=bold_all)


def _highlight(run, color="yellow"):
    rPr = run._r.get_or_add_rPr()
    rPr.append(_oxml("w:highlight", **{"w:val": color}))


# ---------------------------------------------------------------------------
# استایل‌ها (جدول ۴-۱ و ۴-۲ شیوه‌نامه)
# ---------------------------------------------------------------------------
def _style_rpr(style, cfg, fa_size, *, bold=False, mono=False):
    rpr = style.element.get_or_add_rPr()
    _apply_fonts(rpr, cfg, fa_size=fa_size, mono=mono)
    if bold:
        rpr.append(_oxml("w:b", **{"w:val": "1"}))
        rpr.append(_oxml("w:bCs", **{"w:val": "1"}))
    rpr.append(_oxml("w:rtl", **{"w:val": "0" if mono else "1"}))


def _style_ppr(style, cfg, *, space_before=0, line_spacing=None, align=None, bidi=True):
    ppr = style.element.get_or_add_pPr()
    if bidi and ppr.find(qn("w:bidi")) is None:
        ppr.insert(0, _oxml("w:bidi", **{"w:val": "1"}))
    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(0)
    if line_spacing:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line_spacing
    if align is not None:
        pf.alignment = align
    pf.keep_with_next = False


def setup_styles(doc, cfg):
    sz = cfg["sizes_fa"]
    hs = cfg["heading_space_before_pt"]

    # Normal — متن اصلی: B Nazanin 14 / TNR 13، فاصله‌ی خطوط ۱/۵، پیش از بند ۶pt
    normal = doc.styles["Normal"]
    _style_rpr(normal, cfg, sz["body"])
    _style_ppr(normal, cfg,
               space_before=cfg["paragraph_space_before_pt"],
               line_spacing=cfg["line_spacing"],
               align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Heading 1..5 — همگی B Nazanin و Bold
    heads = ((1, sz["h1"], hs["chapter"]), (2, sz["h2"], hs["h2"]),
             (3, sz["h3"], hs["h3"]), (4, sz["h4"], hs["h4"]),
             (5, sz["h5"], hs["h5"]))
    for lvl, size, before in heads:
        st = doc.styles[f"Heading {lvl}"]
        st.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        _style_rpr(st, cfg, size, bold=True)
        _style_ppr(st, cfg, space_before=before, line_spacing=1.0,
                   align=WD_ALIGN_PARAGRAPH.RIGHT)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True

    # عنوان شکل/جدول (Pic Title / Table Title): B Nazanin 13 Bold، وسط‌چین
    for name in ("FigureCaption", "TableCaption"):
        st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = doc.styles["Normal"]
        _style_rpr(st, cfg, sz["caption"], bold=True)
        _style_ppr(st, cfg, space_before=4, line_spacing=1.0,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        st.paragraph_format.space_after = Pt(8)

    # عنوان صفحه‌های مقدماتی — عمداً از سبک Heading نیست تا وارد فهرست
    # عنوان‌ها نشود (فهرست شیوه‌نامه تنها فصل‌ها و بخش‌ها را دربر می‌گیرد).
    st = doc.styles.add_style("FrontTitle", WD_STYLE_TYPE.PARAGRAPH)
    st.base_style = doc.styles["Normal"]
    _style_rpr(st, cfg, sz["h2"], bold=True)
    _style_ppr(st, cfg, space_before=18, line_spacing=1.0,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    st.paragraph_format.space_after = Pt(10)

    st = doc.styles.add_style("FrontSubtitle", WD_STYLE_TYPE.PARAGRAPH)
    st.base_style = doc.styles["Normal"]
    _style_rpr(st, cfg, sz["h3"], bold=True)
    _style_ppr(st, cfg, space_before=hs["h3"], line_spacing=1.0,
               align=WD_ALIGN_PARAGRAPH.RIGHT)
    st.paragraph_format.space_after = Pt(6)

    # متن جدول‌ها (In Table): B Nazanin 13
    st = doc.styles.add_style("InTable", WD_STYLE_TYPE.PARAGRAPH)
    st.base_style = doc.styles["Normal"]
    _style_rpr(st, cfg, sz["in_table"])
    _style_ppr(st, cfg, space_before=2, line_spacing=1.0,
               align=WD_ALIGN_PARAGRAPH.RIGHT)
    st.paragraph_format.space_after = Pt(2)

    # مرجع‌ها (FarsiRef): B Nazanin 13، تک‌فاصله
    st = doc.styles.add_style("RefEntry", WD_STYLE_TYPE.PARAGRAPH)
    st.base_style = doc.styles["Normal"]
    _style_rpr(st, cfg, sz["reference"])
    _style_ppr(st, cfg, space_before=6, line_spacing=1.0,
               align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # کد برنامه‌نویسی (Code): Courier New 10، چپ‌چین
    st = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    st.base_style = doc.styles["Normal"]
    _style_rpr(st, cfg, cfg["fonts"]["code_size_pt"], mono=True)
    _style_ppr(st, cfg, space_before=6, line_spacing=1.0,
               align=WD_ALIGN_PARAGRAPH.LEFT, bidi=False)


# ---------------------------------------------------------------------------
# افزودن انواع بلوک
# ---------------------------------------------------------------------------
def add_heading(doc, text, level, cfg, footnotes=None):
    p = doc.add_heading("", level=level)
    if not has_persian(text):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        style_run(r, cfg, rtl=False, bold=True,
                  size_pt=cfg["sizes_fa"][f"h{min(level,5)}"])
        return p
    set_par_bidi(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_inline(p, text, cfg, base_rtl=True,
               size_pt=cfg["sizes_fa"][f"h{min(level,5)}"], bold_all=True,
               footnotes=footnotes)
    return p


def add_front_title(doc, text, cfg, sub=False, footnotes=None):
    """عنوان صفحه‌های مقدماتی (چکیده، فهرست‌ها، تعهدنامه) — بیرون از فهرست عنوان‌ها."""
    p = doc.add_paragraph(style="FrontSubtitle" if sub else "FrontTitle")
    size = cfg["sizes_fa"]["h3" if sub else "h2"]
    if not has_persian(text):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(p.add_run(text), cfg, rtl=False, bold=True, size_pt=size)
        return p
    set_par_bidi(p)
    add_inline(p, text, cfg, base_rtl=True, size_pt=size, bold_all=True,
               footnotes=footnotes)
    return p


def add_body_paragraph(doc, text, cfg, footnotes=None, quote=False):
    p = doc.add_paragraph()
    set_par_bidi(p)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = cfg["line_spacing"]
    pf.space_before = Pt(cfg["paragraph_space_before_pt"])
    pf.space_after = Pt(0)
    if cfg.get("paragraph_indent_cm"):
        pf.first_line_indent = Cm(cfg["paragraph_indent_cm"])
    if quote:
        pf.right_indent = Cm(0.8)
        pf.left_indent = Cm(0.8)
    add_inline(p, text, cfg, base_rtl=True, footnotes=footnotes)
    return p


def add_english_paragraph(doc, text, cfg):
    """بند تماماً انگلیسی: چپ‌به‌راست، بدون w:bidi."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = cfg["line_spacing"]
    pf.space_before = Pt(cfg["paragraph_space_before_pt"])
    pf.space_after = Pt(0)
    for seg in re.split(r"(\*\*.+?\*\*|`[^`]+`)", text):
        if not seg:
            continue
        if seg.startswith("**"):
            style_run(p.add_run(seg[2:-2]), cfg, rtl=False, bold=True)
        elif seg.startswith("`"):
            style_run(p.add_run(seg[1:-1]), cfg, rtl=False, mono=True)
        else:
            style_run(p.add_run(seg), cfg, rtl=False)
    return p


def add_list_item(doc, text, cfg, numbered=False, footnotes=None):
    style = "List Number" if numbered else "List Bullet"
    try:
        p = doc.add_paragraph(style=style)
    except Exception:
        p = doc.add_paragraph()
    set_par_bidi(p)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = cfg["line_spacing"]
    pf.space_before = Pt(3)
    pf.space_after = Pt(0)
    add_inline(p, text, cfg, base_rtl=True, footnotes=footnotes)
    return p


def add_reference_entry(doc, number, body, cfg, footnotes=None):
    """
    یک مدخل فهرست مرجع‌ها با سبک RefEntry.

    زبانِ مدخل از روی «متنِ مرجع» تعیین می‌شود نه از روی رشته‌ی کامل؛ چون
    شماره‌ی مدخل با ارقام فارسی نوشته می‌شود و اگر مبنای تشخیص قرار گیرد،
    هر مدخل لاتین هم فارسی به‌شمار می‌آید و ایتالیکِ نام نشریه اعمال نمی‌شود.
    """
    size = cfg["sizes_fa"]["reference"]
    p = doc.add_paragraph(style="RefEntry")
    if has_persian(body):
        set_par_bidi(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(p, f"[{number}] {body}", cfg, base_rtl=True,
                   size_pt=size, footnotes=footnotes)
        return p

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_run(p.add_run(f"[{number}] "), cfg, rtl=False, size_pt=size)
    for seg in re.split(r"(\*[^*]+\*)", body):
        if not seg:
            continue
        if seg.startswith("*") and seg.endswith("*") and len(seg) > 2:
            r = style_run(p.add_run(seg[1:-1]), cfg, rtl=False, size_pt=size)
            r.italic = True                       # نام نشریه/کتاب: ایتالیک
        else:
            style_run(p.add_run(seg), cfg, rtl=False, size_pt=size)
    return p


def add_code_block(doc, lines, cfg):
    p = doc.add_paragraph(style="CodeBlock")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.left_indent = Cm(0.3)
    pPr = p._p.get_or_add_pPr()
    pPr.append(_oxml("w:shd", **{"w:val": "clear", "w:color": "auto", "w:fill": "F4F4F4"}))
    _box_paragraph(p, size=4, color="BFBFBF")
    for i, line in enumerate(lines):
        run = p.add_run(line)
        style_run(run, cfg, rtl=False, mono=True)
        if i < len(lines) - 1:
            run.add_break()
    return p


def add_figure(doc, cfg, file, number, caption_text, footnotes=None):
    path = os.path.join(FIG_DIR, file)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    if cfg.get("figure_box"):
        _box_paragraph(p)
    if os.path.exists(path):
        run = p.add_run()
        try:
            run.add_picture(path, width=Cm(14))
        except Exception as e:
            style_run(p.add_run(f"[خطا در درج شکل: {file} — {e}]"), cfg,
                      rtl=True, color=RGBColor(0xC0, 0, 0))
    else:
        r = p.add_run(f"[شکل یافت نشد: figures/{file}]")
        style_run(r, cfg, rtl=True, bold=True, color=RGBColor(0xC0, 0, 0))
        _highlight(r, "yellow")
    # عنوان شکل — زیر شکل (بخش ۴-۹)، قالب «شکل ۴-۱ عنوان»
    cap = doc.add_paragraph(style="FigureCaption")
    set_par_bidi(cap)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep = cfg.get("caption_separator", " ")
    add_inline(cap, f"شکل {number}{sep}{caption_text}", cfg, base_rtl=True,
               size_pt=cfg["sizes_fa"]["caption"], bold_all=True, footnotes=footnotes)
    return p


def add_table_caption(doc, cfg, number, text, footnotes=None):
    cap = doc.add_paragraph(style="TableCaption")
    set_par_bidi(cap)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(10)
    cap.paragraph_format.space_after = Pt(2)
    sep = cfg.get("caption_separator", " ")
    add_inline(cap, f"جدول {number}{sep}{text}", cfg, base_rtl=True,
               size_pt=cfg["sizes_fa"]["caption"], bold_all=True, footnotes=footnotes)


def add_table(doc, rows, cfg, footnotes=None):
    ncol = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncol)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # چیدمان راست‌به‌چپِ ستون‌ها
    tblPr = table._tbl.tblPr
    if tblPr.find(qn("w:bidiVisual")) is None:
        tblPr.append(_oxml("w:bidiVisual", **{"w:val": "1"}))
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            p = cell.paragraphs[0]
            p.text = ""
            p.style = doc.styles["InTable"]
            set_par_bidi(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            add_inline(p, cell_text.strip(), cfg, base_rtl=True,
                       size_pt=cfg["sizes_fa"]["in_table"], bold_all=(i == 0),
                       footnotes=footnotes)
            if i == 0:
                cell._tc.get_or_add_tcPr().append(
                    _oxml("w:shd", **{"w:val": "clear", "w:color": "auto",
                                      "w:fill": "EDEDED"}))
    return table


def add_screenshot_placeholder(doc, cfg, code):
    p = doc.add_paragraph()
    set_par_bidi(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p._p.get_or_add_pPr().append(
        _oxml("w:shd", **{"w:val": "clear", "w:color": "auto", "w:fill": "FFF59D"}))
    pf = p.paragraph_format
    pf.space_before = Pt(6); pf.space_after = Pt(6)
    r = p.add_run(f"[[{code}]]  ← این‌جا اسکرین‌شات درج شود (جزئیات در screenshots.md)")
    style_run(r, cfg, rtl=True, bold=True, color=RGBColor(0xC0, 0x00, 0x00))
    _highlight(r, "yellow")


# ---------------------------------------------------------------------------
# فیلدها (TOC / LOF / LOT / شماره صفحه) و سربرگ/ته‌برگ
# ---------------------------------------------------------------------------
def add_field(paragraph, instr, cfg):
    run = paragraph.add_run()
    run._r.append(_oxml("w:fldChar", **{"w:fldCharType": "begin"}))
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = instr
    run._r.append(instrText)
    run._r.append(_oxml("w:fldChar", **{"w:fldCharType": "separate"}))
    style_run(paragraph.add_run("برای به‌روزرسانی این فهرست: Ctrl+A سپس F9"), cfg, rtl=True)
    run3 = paragraph.add_run()
    run3._r.append(_oxml("w:fldChar", **{"w:fldCharType": "end"}))


def add_toc(doc, cfg, instr, title, footnotes=None):
    add_front_title(doc, title, cfg, footnotes=footnotes)
    p = doc.add_paragraph()
    set_par_bidi(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_field(p, instr, cfg)


def add_page_number_footer(section, cfg):
    """شماره‌ی صفحه در پایین و وسط (بخش ۵-۱۲)."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    set_par_bidi(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run._r.append(_oxml("w:fldChar", **{"w:fldCharType": "begin"}))
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    run._r.append(it)
    run._r.append(_oxml("w:fldChar", **{"w:fldCharType": "separate"}))
    run._r.append(_oxml("w:fldChar", **{"w:fldCharType": "end"}))
    style_run(run, cfg, rtl=True, size_pt=cfg["sizes_fa"]["caption"])


def set_header(section, cfg, text):
    """عنوان فصل در سمت راست سربرگ + خط پررنگ زیر آن (بخش ۴-۸)."""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    set_par_bidi(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    if text:
        add_inline(p, text, cfg, base_rtl=True, size_pt=cfg["sizes_fa"]["caption"])
        _add_bottom_border(p)


def clear_header(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    for r in list(p.runs):
        r._r.getparent().remove(r._r)


def _set_page_num_format(section, fmt, start=None):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = _oxml("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:fmt"), fmt)
    if start is not None:
        pgNumType.set(qn("w:start"), str(start))
    else:
        # سکشن تازه، sectPr سکشن پیشین را کپی می‌کند؛ اگر w:start ارثی حذف نشود،
        # شماره‌ی صفحه در آغاز هر فصل دوباره از ۱ شروع می‌شود.
        if pgNumType.get(qn("w:start")) is not None:
            del pgNumType.attrib[qn("w:start")]


def apply_page_setup(section, cfg):
    mg = cfg["margins_cm"]
    section.top_margin = Cm(mg["top"])
    section.bottom_margin = Cm(mg["bottom"])
    section.right_margin = Cm(mg["right"])   # لبه‌ی صحافی در سند راست‌به‌چپ
    section.left_margin = Cm(mg["left"])
    section.header_distance = Cm(mg.get("header_cm", 1.5))
    section.footer_distance = Cm(mg.get("footer_cm", 1.5))
    set_section_rtl(section)


# ---------------------------------------------------------------------------
# شماره‌گذاری
# ---------------------------------------------------------------------------
class Numbering:
    """شماره‌گذاری فصل-ترتیب برای عنوان‌ها، شکل‌ها و جدول‌ها (بخش ۴-۹)."""

    def __init__(self):
        self.chapter = None      # int یا "پ" برای پیوست
        self.fig = 0
        self.tbl = 0
        self.h2 = 0
        self.h3 = 0
        self.h4 = 0

    def set_chapter(self, ch):
        self.chapter = ch
        self.fig = self.tbl = self.h2 = self.h3 = self.h4 = 0

    def _prefix(self):
        if self.chapter is None:
            return ""
        if isinstance(self.chapter, str):
            return self.chapter
        return fa_num(self.chapter)

    def next_fig(self):
        self.fig += 1
        return f"{self._prefix()}-{fa_num(self.fig)}"

    def next_tbl(self):
        self.tbl += 1
        return f"{self._prefix()}-{fa_num(self.tbl)}"

    def next_heading(self, level):
        """شماره‌ی عنوان: ۲-۱- ، ۲-۱-۳- و ..."""
        if self.chapter is None:
            return ""
        if level == 2:
            self.h2 += 1; self.h3 = 0; self.h4 = 0
            return f"{self._prefix()}-{fa_num(self.h2)}- "
        if level == 3:
            self.h3 += 1; self.h4 = 0
            return f"{self._prefix()}-{fa_num(self.h2)}-{fa_num(self.h3)}- "
        if level == 4:
            self.h4 += 1
            return f"{self._prefix()}-{fa_num(self.h2)}-{fa_num(self.h3)}-{fa_num(self.h4)}- "
        return ""


def _parse_kv(s):
    out = {}
    for part in s.split("|"):
        if "=" in part:
            k, v = part.split("=", 1)
            out[k.strip()] = v.strip()
    return out


# ---------------------------------------------------------------------------
# صفحه‌ی جداکننده‌ی فصل
# ---------------------------------------------------------------------------
def build_chapter_opener(doc, cfg, chapter, title, footnotes=None):
    """صفحه‌ی آغاز فصل: شماره‌ی فصل و عنوان آن (مطابق قالب شیوه‌نامه)."""
    for _ in range(6):
        doc.add_paragraph()
    if isinstance(chapter, int):
        p = doc.add_paragraph()
        set_par_bidi(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(12)
        word = CHAPTER_WORDS[chapter] if chapter < len(CHAPTER_WORDS) else fa_num(chapter)
        style_run(p.add_run(f"فصل {word}"), cfg, rtl=True, bold=True,
                  size_pt=cfg["sizes_fa"]["h2"])
        title = f"{fa_num(chapter)}- {title}"
    h = add_heading(doc, title, 1, cfg, footnotes)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h.paragraph_format.space_before = Pt(cfg["heading_space_before_pt"]["chapter"])
    doc.add_page_break()


# ---------------------------------------------------------------------------
# پارس یک فایل مارک‌داون → افزودن به سند
# ---------------------------------------------------------------------------
def render_markdown(doc, cfg, text, numbering, stats, footnotes, *,
                    chapter_opener=False, chapter=None, references=False,
                    front=False):
    lines = text.split("\n")
    i = 0
    pending_table_caption = None
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("<!--"):
            while i < len(lines) and "-->" not in lines[i]:
                i += 1
            i += 1
            continue

        if stripped == "{{PAGEBREAK}}":
            doc.add_page_break(); i += 1; continue

        if not stripped:
            i += 1; continue

        # عنوان فصل
        if stripped.startswith("# "):
            title = stripped[2:].strip()
            if chapter_opener:
                build_chapter_opener(doc, cfg, chapter, title, footnotes)
            elif front:
                add_front_title(doc, title, cfg, footnotes=footnotes)
            else:
                add_heading(doc, title, 1, cfg, footnotes)
            i += 1; continue

        # عنوان‌های سطح ۲..۴ با شماره‌گذاری خودکار
        m = re.match(r"^(#{2,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            title = m.group(2).strip()
            if front:
                add_front_title(doc, title, cfg, sub=True, footnotes=footnotes)
            else:
                prefix = numbering.next_heading(level)
                add_heading(doc, f"{prefix}{title}", level, cfg, footnotes)
            i += 1; continue

        # بلوک کد
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i]); i += 1
            i += 1
            # اگر عنوانی برای این فهرستِ کد ثبت شده، بالای آن درج می‌شود
            if pending_table_caption:
                add_table_caption(doc, cfg, *pending_table_caption, footnotes=footnotes)
                pending_table_caption = None
            add_code_block(doc, code_lines, cfg)
            stats["code_blocks"] += 1
            continue

        # شکل
        if stripped.startswith("[FIGURE") and stripped.endswith("]"):
            kv = _parse_kv(stripped[7:-1])
            num = numbering.next_fig()
            add_figure(doc, cfg, kv.get("file", "MISSING.png"), num,
                       kv.get("caption", ""), footnotes)
            stats["figures"].append((num, kv.get("file", ""), kv.get("caption", "")))
            i += 1; continue

        # عنوان جدول (برای جدولِ بعدی)
        if stripped.startswith("[TABLE") and stripped.endswith("]"):
            kv = _parse_kv(stripped[6:-1])
            num = numbering.next_tbl()
            pending_table_caption = (num, kv.get("caption", ""))
            stats["tables"].append((num, kv.get("caption", "")))
            i += 1; continue

        mm = re.fullmatch(r"\[\[SCREENSHOT-([0-9A-Za-z]+)\]\]", stripped)
        if mm:
            add_screenshot_placeholder(doc, cfg, f"SCREENSHOT-{mm.group(1)}")
            stats["screenshots"].add(f"SCREENSHOT-{mm.group(1)}")
            i += 1; continue

        # جدول
        if stripped.startswith("|") and stripped.count("|") >= 2:
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i].strip()); i += 1
            rows = []
            for bl in block:
                if set(bl) <= set("|-: "):
                    continue
                rows.append([c.strip() for c in bl.strip("|").split("|")])
            if rows:
                ncol = len(rows[0])
                rows = [(r + [""] * ncol)[:ncol] for r in rows]
                if pending_table_caption:
                    add_table_caption(doc, cfg, *pending_table_caption, footnotes=footnotes)
                    pending_table_caption = None
                add_table(doc, rows, cfg, footnotes)
                sp = doc.add_paragraph()
                sp.paragraph_format.space_after = Pt(0)
                sp.paragraph_format.space_before = Pt(0)
            continue

        # نقل‌قول / تذکر
        if stripped.startswith("> "):
            add_body_paragraph(doc, stripped[2:].strip(), cfg, footnotes, quote=True)
            i += 1; continue

        # فهرست نقطه‌ای
        if stripped.startswith("- ") or stripped.startswith("• "):
            add_list_item(doc, stripped[2:].strip(), cfg, numbered=False,
                          footnotes=footnotes)
            i += 1; continue

        # فهرست عددی / مدخل مرجع
        m = re.match(r"^\d+\.\s+(.*)$", stripped)
        if m:
            body = m.group(1).strip()
            if references:
                idx = int(re.match(r"^(\d+)\.", stripped).group(1))
                add_reference_entry(doc, fa_num(idx), body, cfg, footnotes)
            else:
                add_list_item(doc, body, cfg, numbered=True, footnotes=footnotes)
            i += 1; continue

        # بند معمولی
        para = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not _is_block_start(lines[i].strip()):
            para.append(lines[i].strip()); i += 1
        joined = " ".join(para)
        if has_persian(joined):
            add_body_paragraph(doc, joined, cfg, footnotes)
        else:
            add_english_paragraph(doc, joined, cfg)

    for m in re.finditer(r"\[\[(REF|DATA)-[0-9A-Za-z]+:[^\]]*\]\]", text):
        stats["refs_data"].add(m.group(0))


def _is_block_start(s):
    if s.startswith(("#", "```", "|", "- ", "• ", "> ", "[FIGURE", "[TABLE", "{{")):
        return True
    if re.match(r"^\d+\.\s+", s):
        return True
    if re.fullmatch(r"\[\[SCREENSHOT-[0-9A-Za-z]+\]\]", s):
        return True
    return False


# ---------------------------------------------------------------------------
# صفحه‌های عنوان
# ---------------------------------------------------------------------------
def _center_line(doc, cfg, text, size, bold=False, rtl=True, space_after=6,
                 space_before=0):
    p = doc.add_paragraph()
    if rtl:
        set_par_bidi(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style_run(p.add_run(text), cfg, rtl=rtl, bold=bold, size_pt=size)
    return p


def build_title_page(doc, cfg):
    """صفحه‌ی عنوان فارسی — چیدمان قالب شیوه‌نامه."""
    m = cfg["meta"]
    _center_line(doc, cfg, m.get("university_fa", ""), 18, bold=True, space_before=24)
    _center_line(doc, cfg, m.get("university_sub_fa", ""), 16, bold=True)
    _center_line(doc, cfg, m.get("faculty_fa", ""), 16, bold=True, space_after=30)
    _center_line(doc, cfg, f"پایان‌نامه‌ی دوره‌ی {m.get('degree_fa','کارشناسی')}", 14,
                 space_after=2)
    _center_line(doc, cfg, f"رشته‌ی {m.get('major_fa','')}", 14, space_after=30)
    _center_line(doc, cfg, m.get("title_fa", ""), 18, bold=True, space_after=34)
    _center_line(doc, cfg, "نگارش", 14, space_after=4)
    _center_line(doc, cfg, m.get("student_name_fa", ""), 16, bold=True, space_after=2)
    _center_line(doc, cfg, f"شماره‌ی دانشجویی: {m.get('student_id','')}", 13,
                 space_after=26)
    _center_line(doc, cfg, "استاد راهنما", 14, space_after=4)
    _center_line(doc, cfg, m.get("supervisor_fa", ""), 16, bold=True, space_after=34)
    _center_line(doc, cfg, m.get("defense_term_fa", ""), 14, bold=True)
    doc.add_page_break()


def build_en_title_page(doc, cfg):
    """صفحه‌ی عنوان انگلیسی — واپسین صفحه‌ی گزارش."""
    m = cfg["meta"]
    _center_line(doc, cfg, m.get("university_en", ""), 18, bold=True, rtl=False,
                 space_before=24)
    _center_line(doc, cfg, m.get("university_sub_en", ""), 16, bold=True, rtl=False)
    _center_line(doc, cfg, m.get("faculty_en", ""), 16, bold=True, rtl=False,
                 space_after=30)
    _center_line(doc, cfg, "B.Sc. Thesis", 14, rtl=False, space_after=30)
    _center_line(doc, cfg, m.get("title_en", ""), 18, bold=True, rtl=False,
                 space_after=34)
    _center_line(doc, cfg, "By", 14, rtl=False, space_after=4)
    _center_line(doc, cfg, m.get("student_name_en", ""), 16, bold=True, rtl=False,
                 space_after=26)
    _center_line(doc, cfg, "Supervisor", 14, rtl=False, space_after=4)
    _center_line(doc, cfg, m.get("supervisor_en", ""), 16, bold=True, rtl=False,
                 space_after=34)
    _center_line(doc, cfg, m.get("defense_term_en", ""), 14, bold=True, rtl=False)


# ---------------------------------------------------------------------------
# مانیفست سند
# ---------------------------------------------------------------------------
# kind:
#   gen:title / gen:en_title / gen:toc / gen:lof / gen:lot
#   sec:front-numbered  → شروع شماره‌گذاری أ، ب، ت ...
#   md   → فایل محتوایی (front matter)
#   chapter → فصل (سکشن مستقل + سربرگ + صفحه‌ی جداکننده)
#   back  → مرجع‌ها/پیوست‌ها/چکیده‌ی انگلیسی (سکشن مستقل + سربرگ، بدون شماره‌ی فصل)
MANIFEST = [
    ("gen:title", {}),
    ("md", {"file": "00a-approval.md"}),
    ("sec:front-numbered", {}),
    ("md", {"file": "00b-acknowledgment.md"}),
    ("md", {"file": "00c-abstract-fa.md"}),
    ("gen:toc", {}),
    ("gen:lof", {}),
    ("gen:lot", {}),
    ("md", {"file": "00d-symbols.md"}),
    ("chapter", {"file": "01-intro.md", "ch": 1, "header": "فصل یک: مقدمه", "first": True}),
    ("chapter", {"file": "02-background.md", "ch": 2,
                 "header": "فصل دو: مبانی نظری و کارهای مرتبط"}),
    ("chapter", {"file": "03-analysis-design.md", "ch": 3,
                 "header": "فصل سه: تحلیل و طراحی"}),
    ("chapter", {"file": "04-implementation.md", "ch": 4,
                 "header": "فصل چهار: پیاده‌سازی"}),
    ("chapter", {"file": "05-evaluation.md", "ch": 5,
                 "header": "فصل پنج: ارزیابی و آزمون"}),
    ("chapter", {"file": "06-conclusion.md", "ch": 6,
                 "header": "فصل شش: جمع‌بندی، نتیجه‌گیری و پیشنهادها"}),
    ("back", {"file": "07-references.md", "header": "مرجع‌ها", "references": True}),
    ("back", {"file": "08-appendix.md", "header": "پیوست‌ها", "ch": "پ"}),
    ("back", {"file": "09-abstract-en.md", "header": "Abstract", "front": True}),
    ("gen:en_title", {}),
]


def _new_section(doc, cfg):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_page_setup(sec, cfg)
    return sec


def build():
    cfg = load_cfg()
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = Document()

    sec = doc.sections[0]
    apply_page_setup(sec, cfg)
    setup_styles(doc, cfg)

    footnotes = Footnotes(cfg)
    numbering = Numbering()
    stats = {"figures": [], "tables": [], "screenshots": set(), "refs_data": set(),
             "code_blocks": 0, "missing_files": [], "chapters": []}

    # سکشن نخست: صفحه‌ی عنوان و تعهدنامه — بدون سربرگ و بدون شماره‌ی صفحه
    clear_header(sec)
    sec.footer.is_linked_to_previous = False
    body_started = False

    for kind, args in MANIFEST:
        if kind == "gen:title":
            build_title_page(doc, cfg)

        elif kind == "sec:front-numbered":
            s = _new_section(doc, cfg)
            clear_header(s)
            _set_page_num_format(s, cfg["page_numbers"]["front_matter"], start=1)
            add_page_number_footer(s, cfg)

        elif kind == "gen:toc":
            add_toc(doc, cfg, 'TOC \\o "1-3" \\h \\z \\u', cfg["toc"]["title"], footnotes)
            doc.add_page_break()

        elif kind == "gen:lof":
            add_toc(doc, cfg, 'TOC \\h \\z \\t "FigureCaption,1"',
                    cfg["toc"]["figures_title"], footnotes)
            doc.add_page_break()

        elif kind == "gen:lot":
            add_toc(doc, cfg, 'TOC \\h \\z \\t "TableCaption,1"',
                    cfg["toc"]["tables_title"], footnotes)
            doc.add_page_break()

        elif kind == "gen:en_title":
            s = _new_section(doc, cfg)
            clear_header(s)
            build_en_title_page(doc, cfg)

        elif kind in ("md", "chapter", "back"):
            path = os.path.join(CONTENT_DIR, args["file"])
            if not os.path.exists(path):
                stats["missing_files"].append(args["file"]); continue
            with open(path, encoding="utf-8") as f:
                text = f.read()

            if kind == "md":
                numbering.set_chapter(None)
                render_markdown(doc, cfg, text, numbering, stats, footnotes,
                                front=True)
                doc.add_page_break()
            else:
                s = _new_section(doc, cfg)
                set_header(s, cfg, args.get("header", ""))
                if not body_started:
                    # آغاز متن اصلی: شماره‌گذاری از ۱ با ارقام لاتین/فارسی
                    _set_page_num_format(s, cfg["page_numbers"]["body"], start=1)
                    body_started = True
                else:
                    _set_page_num_format(s, cfg["page_numbers"]["body"])
                add_page_number_footer(s, cfg)
                numbering.set_chapter(args.get("ch"))
                render_markdown(doc, cfg, text, numbering, stats, footnotes,
                                chapter_opener=(kind == "chapter"),
                                chapter=args.get("ch"),
                                references=args.get("references", False),
                                front=args.get("front", False))
                if kind == "chapter":
                    stats["chapters"].append((args["ch"], args.get("header", "")))

    footnotes.attach(doc)
    doc.save(OUT_FILE)
    validate(cfg, stats, footnotes)


# ---------------------------------------------------------------------------
# اعتبارسنجی خودکار
# ---------------------------------------------------------------------------
def validate(cfg, stats, footnotes):
    print("\n" + "=" * 64)
    print(" گزارش ساخت سند ".center(64, "="))
    print("=" * 64)
    print(f"خروجی: {OUT_FILE}")
    print(f"فصل‌ها: {len(stats['chapters'])}   شکل‌ها: {len(stats['figures'])}   "
          f"جدول‌ها: {len(stats['tables'])}   بلوک‌کد: {stats['code_blocks']}   "
          f"پاورقی: {len(footnotes.items)}   اسکرین‌شات: {len(stats['screenshots'])}")

    if stats["missing_files"]:
        print("\n⚠️  فایل‌های محتوایی یافت‌نشده:")
        for f in stats["missing_files"]:
            print(f"     - content/{f}")
    else:
        print("\n✅ همه‌ی فایل‌های محتوایی موجود بودند.")

    missing_figs = [(n, f) for n, f, _ in stats["figures"]
                    if f and not os.path.exists(os.path.join(FIG_DIR, f))]
    if missing_figs:
        print("\n⚠️  شکل‌های بدون فایل:")
        for num, file in missing_figs:
            print(f"     - شکل {num}: {file}")
    else:
        print("✅ همه‌ی شکل‌های ارجاع‌شده فایل دارند.")

    print("\n— فهرست شکل‌ها:")
    for num, file, cap in stats["figures"]:
        print(f"     شکل {num} {cap}  ({file})")
    print("— فهرست جدول‌ها:")
    for num, cap in stats["tables"]:
        print(f"     جدول {num} {cap}")

    if stats["refs_data"]:
        print("\n⚠️  نشانه‌های حل‌نشده‌ی REF/DATA در متن:")
        for tok in sorted(stats["refs_data"]):
            print(f"     - {tok[:90]}")
    else:
        print("\n✅ هیچ نشانه‌ی REF/DATA حل‌نشده‌ای در متن نمانده است.")

    if stats["screenshots"]:
        print(f"\nℹ️  {len(stats['screenshots'])} جای اسکرین‌شات در متن هست "
              f"(فهرست و راهنمای تهیه: thesis/screenshots.md).")

    print("\nیادآوری: در Word کلید Ctrl+A و سپس F9 را بزنید تا فهرست عنوان‌ها،"
          " فهرست شکل‌ها/جدول‌ها و شماره‌ی صفحه‌ها به‌روز شوند.")
    print("=" * 64 + "\n")


if __name__ == "__main__":
    build()
