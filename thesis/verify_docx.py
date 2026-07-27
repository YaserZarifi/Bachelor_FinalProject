# -*- coding: utf-8 -*-
"""
verify_docx.py — راستی‌آزمایی خروجی report.docx در برابر شیوه‌نامه‌ی امیرکبیر

دو گروه بررسی انجام می‌شود:
  ۱) نامتغیرهای راست‌به‌چپ (bidi / rtl / szCs)
  ۲) الزام‌های عددی شیوه‌نامه: حاشیه‌ها، فاصله‌ی خطوط، قلم‌ها، سربرگ و ته‌برگ،
     شماره‌گذاری صفحه‌های مقدماتی و وجود پاورقی‌ها

نکته‌ی مهم: پس از آن‌که سند یک بار در Word باز و ذخیره شود، Word قالب‌بندیِ
مستقیمِ تکراری را حذف و آن را به سبک‌ها منتقل می‌کند. بنابراین این اسکریپت
قالب‌بندیِ «مؤثر» را می‌سنجد، یعنی زنجیره‌ی سبک‌ها را هم دنبال می‌کند.

اجرا:  python thesis/verify_docx.py
"""
import os
import re
import zipfile

from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
DOC = os.path.join(HERE, "out", "report.docx")
PERSIAN = re.compile(r"[؀-ۿ]")


def has_fa(t):
    return bool(PERSIAN.search(t or ""))


doc = Document(DOC)

# ---------------------------------------------------------------------------
# استخراج آنچه هر سبک (و زنجیره‌ی سبک‌های پایه‌اش) تأمین می‌کند
# ---------------------------------------------------------------------------
def style_provides(style, seen=None):
    """(bidi, szCs, rtl) که این سبک یا سبک‌های پایه‌ی آن تأمین می‌کنند."""
    if style is None:
        return (False, False, False)
    seen = seen or set()
    if style.style_id in seen:
        return (False, False, False)
    seen.add(style.style_id)

    el = style.element
    ppr = el.find(qn("w:pPr"))
    rpr = el.find(qn("w:rPr"))
    bidi = ppr is not None and ppr.find(qn("w:bidi")) is not None
    szcs = rpr is not None and rpr.find(qn("w:szCs")) is not None
    rtl = rpr is not None and rpr.find(qn("w:rtl")) is not None

    base = style_provides(style.base_style, seen)
    return (bidi or base[0], szcs or base[1], rtl or base[2])


STYLE_CACHE = {}


def provided(style):
    name = getattr(style, "name", None) or "?"
    if name not in STYLE_CACHE:
        STYLE_CACHE[name] = style_provides(style)
    return STYLE_CACHE[name]


# ---------------------------------------------------------------------------
# گروه ۱ — نامتغیرهای راست‌به‌چپ
# ---------------------------------------------------------------------------
stats = dict(paras=0, fa_paras=0, no_bidi=0, fa_runs=0, no_szcs=0, no_rtl=0)
bad = []

for p in doc.paragraphs:
    stats["paras"] += 1
    style = p.style
    sname = style.name or ""
    st_bidi, st_szcs, st_rtl = provided(style)
    if has_fa(p.text) and sname != "CodeBlock":
        stats["fa_paras"] += 1
        pPr = p._p.find(qn("w:pPr"))
        own = pPr is not None and pPr.find(qn("w:bidi")) is not None
        if not (own or st_bidi):
            stats["no_bidi"] += 1
            bad.append(f"[{sname}] {p.text[:60]}")
    for r in p.runs:
        if not has_fa(r.text):
            continue
        stats["fa_runs"] += 1
        rPr = r._r.find(qn("w:rPr"))
        own_szcs = rPr is not None and rPr.find(qn("w:szCs")) is not None
        own_rtl = rPr is not None and rPr.find(qn("w:rtl")) is not None
        if not (own_szcs or st_szcs):
            stats["no_szcs"] += 1
        if not (own_rtl or st_rtl):
            stats["no_rtl"] += 1

# ---------------------------------------------------------------------------
# گروه ۲ — الزام‌های عددی شیوه‌نامه
# ---------------------------------------------------------------------------
EXPECT = {"top": 3.0, "bottom": 3.0, "right": 3.0, "left": 2.5}
margin_problems = []
for i, sec in enumerate(doc.sections):
    got = {
        "top": round(sec.top_margin.cm, 2),
        "bottom": round(sec.bottom_margin.cm, 2),
        "right": round(sec.right_margin.cm, 2),
        "left": round(sec.left_margin.cm, 2),
    }
    for k, v in EXPECT.items():
        if abs(got[k] - v) > 0.05:
            margin_problems.append(f"سکشن {i}: {k}={got[k]} (انتظار {v})")

normal = doc.styles["Normal"]
line_spacing = normal.paragraph_format.line_spacing
space_before = normal.paragraph_format.space_before

def style_font(name):
    """قلم و اندازه‌ی مؤثر یک سبک؛ زنجیره‌ی basedOn تا Normal دنبال می‌شود."""
    st = doc.styles[name]
    cs = size = None
    seen = set()
    while st is not None and st.style_id not in seen:
        seen.add(st.style_id)
        rpr = st.element.find(qn("w:rPr"))
        if rpr is not None:
            rf = rpr.find(qn("w:rFonts"))
            if cs is None and rf is not None:
                # w:cstheme بر w:cs اولویت دارد؛ اگر تم باشد یعنی قلم دلخواه اعمال نشده
                cs = rf.get(qn("w:cstheme")) or rf.get(qn("w:cs"))
            szcs = rpr.find(qn("w:szCs"))
            if size is None and szcs is not None:
                size = int(szcs.get(qn("w:val"))) / 2
        st = st.base_style
    return (cs or "?", size if size is not None else "?")

# سربرگ و ته‌برگ
headers = sum(1 for s in doc.sections if s.header.paragraphs[0].text.strip())
footers_with_page = 0
for s in doc.sections:
    if "PAGE" in s.footer.paragraphs[0]._p.xml:
        footers_with_page += 1

# شماره‌گذاری صفحه‌های مقدماتی و متن اصلی
fmts = []
for s in doc.sections:
    pg = s._sectPr.find(qn("w:pgNumType"))
    if pg is not None:
        fmts.append((pg.get(qn("w:fmt")), pg.get(qn("w:start"))))

# پاورقی‌ها
with zipfile.ZipFile(DOC) as z:
    names = z.namelist()
    fn_part = "word/footnotes.xml" in names
    fn_count = 0
    if fn_part:
        blob = z.read("word/footnotes.xml").decode("utf-8")
        fn_count = len(re.findall(r'<w:footnote w:id="[1-9]', blob))

# ---------------------------------------------------------------------------
# گزارش
# ---------------------------------------------------------------------------
print("=" * 62)
print(" راستی‌آزمایی خروجی در برابر شیوه‌نامه ".center(62, "="))
print("=" * 62)
print(f"پاراگراف‌ها: {stats['paras']}  (فارسی: {stats['fa_paras']})   "
      f"جدول: {len(doc.tables)}   تصویر: {len(doc.inline_shapes)}")
print(f"سکشن‌ها: {len(doc.sections)}   سربرگِ دارای متن: {headers}   "
      f"ته‌برگِ دارای شماره‌ی صفحه: {footers_with_page}")
print("-" * 62)

ok = True


def check(label, condition, detail=""):
    global ok
    if not condition:
        ok = False
    print(f"{'✅' if condition else '❌'} {label}" + (f"  ({detail})" if detail else ""))


check("پاراگراف فارسی بدون جهت راست‌به‌چپ مؤثر", stats["no_bidi"] == 0,
      f"{stats['no_bidi']} مورد")
check("run فارسی بدون اندازه‌ی Complex Script مؤثر", stats["no_szcs"] == 0,
      f"{stats['no_szcs']} مورد")
check("run فارسی بدون علامت rtl مؤثر", stats["no_rtl"] == 0,
      f"{stats['no_rtl']} مورد")
check("حاشیه‌ها ۳/۳/۳/۲٫۵ سانتی‌متر (بخش ۴-۴-۱)", not margin_problems,
      "؛ ".join(margin_problems[:3]))
check("فاصله‌ی خطوط ۱٫۵ روی سبک Normal (بخش ۴-۴-۲)",
      line_spacing is not None and abs(line_spacing - 1.5) < 0.01, str(line_spacing))
check("فاصله‌ی ۶pt پیش از هر بند (بخش ۴-۴-۳)",
      space_before is not None and abs(space_before.pt - 6) < 0.01,
      f"{space_before.pt if space_before else '?'}pt")

for sname, exp_font, exp_size in [("Normal", "B Nazanin", 14),
                                  ("Heading 1", "B Nazanin", 20),
                                  ("Heading 2", "B Nazanin", 18),
                                  ("Heading 3", "B Nazanin", 16),
                                  ("FigureCaption", "B Nazanin", 13),
                                  ("TableCaption", "B Nazanin", 13)]:
    font, size = style_font(sname)
    check(f"قلم و اندازه‌ی سبک {sname} = {exp_font} {exp_size} (جدول ۴-۱)",
          font == exp_font and size == exp_size, f"{font} {size}")

check("سربرگ با عنوان فصل روی همه‌ی سکشن‌های متن اصلی", headers >= 6,
      f"{headers} سکشن")
check("شماره‌ی صفحه در ته‌برگ", footers_with_page >= 7, f"{footers_with_page} سکشن")
check("شماره‌گذاری مقدماتی به‌صورت أ، ب، ت (arabicAlpha)",
      any(f == "arabicAlpha" for f, _ in fmts))
check("شماره‌گذاری متن اصلی از ۱ (decimal پیش‌فرضِ حذف‌شده توسط Word هم پذیرفته است)",
      any(f in (None, "decimal") and s == "1" for f, s in fmts))
check("بخش پاورقی‌ها در بسته موجود است (بخش ۳-۱)", fn_part and fn_count > 0,
      f"{fn_count} پاورقی")

print("-" * 62)
if bad:
    print("نمونه‌ی پاراگراف‌های مشکل‌دار:")
    for b in bad[:8]:
        print("   ·", b)
    print("-" * 62)
print("نتیجه:", "همه‌ی بررسی‌ها موفق ✅" if ok else "نقض یافت شد ❌")
print("=" * 62)
